#!/usr/bin/env python3
"""Simple text extractor for PDFs and images.

Usage:
  python extractor.py /path/to/file_or_dir

Features:
- If a PDF contains extractable text, use pdfminer.six.
- Otherwise convert PDF pages to images (pdf2image) and run OCR (pytesseract).
- For image files, run OCR via pytesseract.
"""
import argparse
import io
import os
import re
import subprocess
import sys
import zipfile
from collections import Counter
from pathlib import Path
from typing import List, Optional

try:
    from pdfminer.high_level import extract_text as pdf_extract_text
except Exception:
    pdf_extract_text = None

try:
    from pdf2image import convert_from_path
except Exception:
    convert_from_path = None

try:
    from PIL import Image, ImageFilter, ImageOps, ImageStat
except Exception:
    Image = None
    ImageFilter = None
    ImageOps = None
    ImageStat = None

try:
    import pytesseract
except Exception:
    pytesseract = None

if pytesseract is not None:
    _tesseract_cmd = os.getenv("TESSERACT_CMD")
    if _tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = _tesseract_cmd
    elif sys.platform == "win32":
        # The Windows installer doesn't add tesseract.exe to PATH by
        # default; fall back to its default install location if present.
        _default_tesseract = Path(os.environ.get("PROGRAMFILES", r"C:\Program Files")) / "Tesseract-OCR" / "tesseract.exe"
        if _default_tesseract.exists():
            pytesseract.pytesseract.tesseract_cmd = str(_default_tesseract)

try:
    import pymupdf
except Exception:
    pymupdf = None

try:
    import cv2
    import numpy as np
except Exception:
    cv2 = None
    np = None

try:
    # Optional: a trained document-layout model, used only as a
    # pre-filter (see _mask_layout_noise below) to mask out figure/table
    # regions before OCR. Not in requirements.txt — see
    # requirements-doclayout.txt and the README for why (AGPL-3.0
    # license, heavy dependency footprint).
    from doclayout_yolo import YOLOv10
except Exception:
    YOLOv10 = None

try:
    # Optional: CRAFT text-region detector (vendored under craft_detector/,
    # MIT License — see craft_detector/LICENSE), used only as a targeted
    # gap-filler (see _craft_fill_gaps below) for content the normal OCR
    # pipeline's own detections miss entirely. Not in requirements.txt —
    # see requirements-craft.txt and the README.
    from craft_detector.detect import detect_boxes as _craft_detect_boxes
    from craft_detector.detect import load_model as _craft_load_model
except Exception:
    _craft_detect_boxes = None
    _craft_load_model = None

try:
    import openpyxl
except Exception:
    openpyxl = None

try:
    import docx as python_docx
except Exception:
    python_docx = None


IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".gif", ".webp"}
TEXT_EXTS = {".txt", ".csv"}
SPREADSHEET_EXTS = {".xlsx", ".xlsm"}
DOCX_EXTS = {".docx"}
OFFICE_ZIP_EXTS = SPREADSHEET_EXTS | DOCX_EXTS

try:
    # Adds HEIC/HEIF decoding support to Pillow (common iOS/macOS screenshot
    # and photo format that Pillow can't open on its own).
    import pillow_heif

    pillow_heif.register_heif_opener()
    IMAGE_EXTS |= {".heic", ".heif"}
except Exception:
    pass

# Tesseract language packs to run together (Armenian docs are often mixed
# with Russian/English headers and stamps, like the NAIRI medical reports).
# "hye" is Tesseract's standard Eastern Armenian model — the modern,
# Yerevan-standard written orthography — which is what actual source
# material here uses; it isn't tuned for Western/Classical Armenian.
# Override with the OCR_LANGS env var, e.g. "eng" for Latin-only scans.
DEFAULT_OCR_LANGS = os.getenv("OCR_LANGS", "hye+rus+eng")

# On Windows, poppler (pdftoppm/pdftocairo) usually isn't on PATH unless
# manually added; point pdf2image at its bin/ folder via this env var.
POPPLER_PATH = os.getenv("POPPLER_PATH") or None


def _installed_ocr_langs() -> Optional[set]:
    """List Tesseract's installed language/model codes.

    Deliberately doesn't use pytesseract.get_languages(): it filters
    tesseract's own `--list-langs` output through a strict `^[a-z_]+$`
    regex that rejects any code with a hyphen or digit — which silently
    drops legitimate custom/community models with such names, even
    though tesseract itself lists and runs them fine. That would make
    _resolve_ocr_langs below wrongly report a real, working language
    pack as missing.
    """
    if pytesseract is None:
        return None
    try:
        result = subprocess.run(
            [pytesseract.pytesseract.tesseract_cmd, "--list-langs"],
            capture_output=True,
            text=True,
            check=False,
        )
    except OSError:
        return None
    # First line is a header ("List of available languages in ..."), the
    # rest is one language/model code per line.
    lines = [line.strip() for line in result.stdout.splitlines() if line.strip()]
    langs = set(lines[1:]) if lines and lines[0].lower().startswith("list of") else set(lines)
    return langs or None


def _resolve_ocr_langs(lang: str) -> str:
    """Validate the '+'-joined language codes against what Tesseract
    actually has trained data for.

    A missing language pack (most commonly Armenian on Windows, where the
    installer's language checkboxes are easy to miss) doesn't reliably
    make tesseract error out — depending on the build, it can silently OCR
    the script using only the languages it does have, misreading Armenian
    letter shapes as vaguely similar Latin/Cyrillic ones and returning
    fluent-looking but meaningless gibberish instead of failing. Checking
    up front turns that into a clear, actionable error.
    """
    requested = [code for code in lang.split("+") if code]
    installed = _installed_ocr_langs()
    if not installed:
        # get_languages() itself failed (e.g. tesseract not found at all);
        # let the OCR call raise its own, more specific error.
        return lang

    missing = [code for code in requested if code not in installed]
    if not missing:
        return lang

    usable = [code for code in requested if code in installed]
    fix_hint = (
        "Install the missing Tesseract language data (see the README's per-OS "
        "instructions — on Windows, rerun the Tesseract installer and check the "
        "language you need)"
    )
    if not usable:
        raise RuntimeError(
            f"None of the requested OCR languages ({'+'.join(requested)}) are installed "
            f"in Tesseract ({len(installed)} other language(s) found). {fix_hint}, "
            "or set OCR_LANGS to a language you do have installed."
        )

    raise RuntimeError(
        f"Missing Tesseract language data for: {', '.join(missing)}. OCR would silently "
        "misread that script using only the installed languages instead of failing "
        f"clearly, so refusing to run. {fix_hint}, or set OCR_LANGS={'+'.join(usable)} "
        "to proceed with only what's installed."
    )


def _preprocess_for_ocr(img: "Image.Image") -> "Image.Image":
    """Upscale small images and boost contrast before OCR.

    Tesseract's layout analysis and character recognition are unreliable
    on the low-resolution, low-contrast images typical of screenshots
    (e.g. small white text on a colored chat-bubble background) — it can
    silently drop whole text blocks rather than misread them.
    """
    img = img.convert("RGB")
    if max(img.size) < 2000:
        img = img.resize((img.width * 2, img.height * 2), Image.LANCZOS)
    gray = ImageOps.autocontrast(ImageOps.grayscale(img))
    # Decorative/script fonts common on posters and graphics are often
    # anti-aliased into soft edges that blur together at OCR resolution;
    # a light unsharp mask crisps the strokes back up before thresholding.
    return gray.filter(ImageFilter.UnsharpMask(radius=2, percent=150, threshold=3))


_WORD_RE = re.compile(r"[^\W\d_]{3,}", re.UNICODE)

# A single character repeated 3+ times in a row (e.g. "AAAAAAA") is
# essentially never real text. It's the classic misread of a repeating
# decorative graphic — a row of pine-tree triangles, dots, snowflakes, a
# striped ribbon edge — into the one Latin/Cyrillic/etc. letter that
# shape resembles. Tesseract is often unnervingly *confident* about that
# misread since the shape repeats cleanly, so left unfiltered it can
# rack up a higher confidence-weighted score than the real (shorter,
# less certain) text elsewhere in a busy image and win the "best of"
# comparison outright.
_REPEATED_CHAR_RE = re.compile(r"(.)\1{2,}")

_HAS_LETTER_RE = re.compile(r"[^\W\d_]", re.UNICODE)


def _looks_like_text(word: str) -> bool:
    """A bare stroke, edge, or blob in a graphic/illustration routinely
    gets misread as a single digit or punctuation mark ("1", "|", ".",
    ":") — plausible-looking, sometimes even high-confidence, but not
    text. Requiring at least one real letter (or, for legitimate
    all-digit content like a page number or a year, at least 2
    characters) filters those isolated fragments out without rejecting
    genuine short numbers.
    """
    if _HAS_LETTER_RE.search(word):
        return True
    return len(word) >= 2


# Rough per-script letter ranges, just enough to tell Armenian/Cyrillic/Latin
# apart for the stray-glyph filter below — not a full script classifier.
_SCRIPT_RANGES = {
    "armenian": re.compile(r"[԰-֏]"),
    "cyrillic": re.compile(r"[Ѐ-ӿ]"),
    "latin": re.compile(r"[A-Za-zÀ-ɏ]"),
}


def _dominant_script(word: str) -> Optional[str]:
    counts = {name: len(pat.findall(word)) for name, pat in _SCRIPT_RANGES.items()}
    best = max(counts, key=counts.get)
    return best if counts[best] > 0 else None


def _strip_stray_script_glyphs(text: str) -> str:
    """Drop short (<=2 char) word-tokens whose script doesn't match the
    rest of their line.

    A single misread glyph — a decorative flourish, a border/divider
    element, noise near a color transition — can land on an otherwise
    clean line of real text and pass every other filter (it has a real
    letter). When that stray glyph happens to be recognized in a
    *different* script than every real word around it (e.g. one Armenian
    letter tacked onto an English caption because the OCR language pack
    covers both), that mismatch is itself strong evidence it's not real
    content — genuine short loanwords in a foreign script are far rarer
    than this specific misread pattern. Longer tokens carry enough of
    their own evidence to stand on their own and are left untouched.

    Applied as a final text-level pass (not inside the word-box
    reconstruction) so it works regardless of which OCR reconstruction
    produced the text — Tesseract's own full-page serialization, or this
    module's word-box reconstruction.
    """
    out_lines = []
    for line in text.split("\n"):
        tokens = line.split(" ")
        scripts = [_dominant_script(t) for t in tokens if len(t) >= 3]
        scripts = [s for s in scripts if s]
        if not scripts:
            out_lines.append(line)
            continue
        dominant = Counter(scripts).most_common(1)[0][0]
        kept = [t for t in tokens if len(t) > 2 or _dominant_script(t) in (dominant, None)]
        out_lines.append(" ".join(kept))
    return "\n".join(out_lines)


def _box_iou(a: tuple, b: tuple) -> float:
    ax, ay, aw, ah = a
    bx, by, bw, bh = b
    ix0, iy0 = max(ax, bx), max(ay, by)
    ix1, iy1 = min(ax + aw, bx + bw), min(ay + ah, by + bh)
    iw, ih = max(0.0, ix1 - ix0), max(0.0, iy1 - iy0)
    inter = iw * ih
    union = aw * ah + bw * bh - inter
    return inter / union if union > 0 else 0.0


def _boxes_compete(a: tuple, b: tuple) -> bool:
    """Whether two boxes are plausibly two different candidates' readings
    of the *same* word position, not two separate words.

    Plain IoU misses a common case: a small fragment box (a different,
    worse candidate splitting a connected word into pieces) sitting
    almost entirely *inside* a correct whole-word box. Their union is
    dominated by the big box, so IoU stays low even though the fragment
    is clearly not an independent word — it's noise from the same spot.
    Checking horizontal-overlap-relative-to-the-smaller-box on same-line
    pairs catches that containment case too.
    """
    if _box_iou(a, b) > 0.3:
        return True
    ax, ay, aw, ah = a
    bx, by, bw, bh = b
    if abs(ay - by) >= max(ah, bh) * 0.6:
        return False  # not even the same line
    overlap_x = max(0.0, min(ax + aw, bx + bw) - max(ax, bx))
    smaller_w = min(aw, bw)
    return smaller_w > 0 and overlap_x / smaller_w > 0.5


def _box_mostly_within(inner: tuple, outer: tuple) -> bool:
    """Whether most of `inner`'s own area falls inside `outer`.

    Unlike _boxes_compete (built for comparing two same-scale word-line
    boxes, with a "same line" gate that assumes both are line-height),
    this is for checking a small detection against a much taller region —
    a detection near the bottom of a tall region can be legitimately
    inside it while still being hundreds of pixels from the region's own
    top edge, which would fail _boxes_compete's line-height-relative gate.
    """
    ix, iy, iw, ih = inner
    ox, oy, ow, oh = outer
    overlap_x = max(0.0, min(ix + iw, ox + ow) - max(ix, ox))
    overlap_y = max(0.0, min(iy + ih, oy + oh) - max(iy, oy))
    inner_area = max(1e-6, iw * ih)
    return (overlap_x * overlap_y) / inner_area > 0.5


def _otsu_threshold(gray_img: "Image.Image") -> int:
    hist = gray_img.histogram()
    total = sum(hist)
    sum_all = sum(i * h for i, h in enumerate(hist))
    sum_bg = 0.0
    weight_bg = 0
    best_variance = 0.0
    threshold = 127
    for i, h in enumerate(hist):
        weight_bg += h
        if weight_bg == 0:
            continue
        weight_fg = total - weight_bg
        if weight_fg == 0:
            break
        sum_bg += i * h
        mean_bg = sum_bg / weight_bg
        mean_fg = (sum_all - sum_bg) / weight_fg
        variance = weight_bg * weight_fg * (mean_bg - mean_fg) ** 2
        if variance > best_variance:
            best_variance = variance
            threshold = i
    return threshold


def _adaptive_threshold_variants(gray_img: "Image.Image") -> List["Image.Image"]:
    """Locally-thresholded binarizations of a grayscale image via OpenCV.

    Otsu's threshold is global: it picks one cutoff for the whole image, so
    it only works when text/background contrast is roughly uniform
    everywhere. Poster-style graphics with gradients, multi-colored text,
    and illustration behind the words routinely have different local
    contrast in different regions — a word that's legible against a
    lighter corner of the background can vanish into a darker corner
    under one global cutoff. Gaussian-weighted adaptive thresholding picks
    a cutoff per neighborhood instead, so it can pick up text a global
    threshold misses. Returns [] if OpenCV/NumPy aren't installed.
    """
    if cv2 is None or np is None:
        return []
    arr = np.array(gray_img)
    variants = []
    for block_size in (25, 51):
        binarized = cv2.adaptiveThreshold(
            arr, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, block_size, 10
        )
        variants.append(Image.fromarray(binarized))
    return variants


def _cv2_clahe_variant(img: "Image.Image") -> Optional["Image.Image"]:
    """A second, independently-tuned preprocessing pipeline via OpenCV:
    cubic-upscale → auto-invert dark images → CLAHE → bilateral denoise →
    adaptive threshold.

    The PIL-based pipeline above (autocontrast + unsharp mask) is cheap
    and works for most images, but CLAHE equalizes local contrast far
    better than one global autocontrast pass on unevenly-lit poster
    photos, and a bilateral filter smooths illustration/photo noise while
    keeping text edges sharp — cases where the simpler pipeline still
    comes up short. Runs off the original (not the PIL-preprocessed)
    image so the two pipelines stay independent. Returns None if
    OpenCV/NumPy aren't installed.
    """
    if cv2 is None or np is None:
        return None
    gray = cv2.cvtColor(np.array(img.convert("RGB")), cv2.COLOR_RGB2GRAY)
    if max(gray.shape) < 2000:
        gray = cv2.resize(gray, None, fx=2, fy=2, interpolation=cv2.INTER_CUBIC)
    if gray.mean() < 127:
        gray = cv2.bitwise_not(gray)
    clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
    enhanced = clahe.apply(gray)
    denoised = cv2.bilateralFilter(enhanced, d=9, sigmaColor=75, sigmaSpace=75)
    thresh = cv2.adaptiveThreshold(
        denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
    )
    return Image.fromarray(thresh)


def _shading_corrected_variant(img: "Image.Image") -> Optional["Image.Image"]:
    """Illumination/shading-corrected adaptive-threshold binarization.

    CLAHE and Otsu/adaptive thresholding correct *local* contrast, but a
    large-scale gradient or vignette across a poster's background (soft
    lighting, a color fade) can still bias them since it changes the
    baseline brightness a whole region sits on. Dividing the image by a
    heavily morphologically-opened (i.e. large-scale-blurred) version of
    itself cancels that gradient out first, so the adaptive threshold
    that follows only has to deal with the actual text strokes on an
    already-flattened background. Returns None if OpenCV/NumPy aren't
    installed.
    """
    if cv2 is None or np is None:
        return None
    gray = cv2.cvtColor(np.array(img.convert("RGB")), cv2.COLOR_RGB2GRAY)
    if max(gray.shape) < 2000:
        gray = cv2.resize(gray, None, fx=2, fy=2, interpolation=cv2.INTER_CUBIC)
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (15, 15))
    background = cv2.morphologyEx(gray, cv2.MORPH_OPEN, kernel)
    background = np.where(background == 0, 1, background).astype(np.uint8)
    normalized = cv2.divide(gray, background, scale=255)
    thresh = cv2.adaptiveThreshold(
        normalized, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 5
    )
    return Image.fromarray(thresh)


# Engine/page-segmentation combos to try alongside the default. --oem 1
# (LSTM-only) is pinned explicitly rather than left at --oem 3 ("default,
# based on what's available"): the hye/rus/eng traineddata this app
# targets is LSTM-only on modern Tesseract, so the two behave identically
# today, but pinning --oem 1 guarantees that instead of silently changing
# behavior if a legacy-inclusive language pack (still shipped by some
# Linux distros) ever ends up installed. --psm 3 ("fully automatic")
# assumes a page laid out in coherent blocks/columns; poster graphics
# scatter short text fragments around illustrations with no such
# structure, which can make automatic layout analysis miss them entirely.
# --psm 6 ("single uniform block") suits a stacked caption like "March /
# Women's / Day", while --psm 11 ("sparse text") looks for words anywhere
# in the image in no particular order — between the three, most poster
# layouts are covered.
_OCR_CONFIGS = ("--oem 1", "--oem 1 --psm 6", "--oem 1 --psm 11")


def _ocr_candidate_images(raw_img: "Image.Image") -> List[tuple]:
    """Build several different renderings of the same image for OCR to try,
    each paired with its (scale_x, scale_y) relative to raw_img.

    Tesseract is tuned for solid dark text on a light background. Posters,
    dark-mode screenshots, and images with busy illustrated backgrounds
    can confuse its thresholding badly enough that a plain grayscale pass
    reads nothing, or hallucinates junk from the artwork. Two independent
    preprocessing pipelines (PIL autocontrast/unsharp, and OpenCV
    CLAHE/bilateral-denoise), each with plain/inverted/Otsu/adaptive-
    threshold binarization, cover the common failure modes without having
    to guess which one applies ahead of time.
    """
    pil_pre = _preprocess_for_ocr(raw_img)
    sx, sy = pil_pre.width / raw_img.width, pil_pre.height / raw_img.height
    threshold = _otsu_threshold(pil_pre)
    binarized = pil_pre.point(lambda p: 255 if p > threshold else 0)
    # Every candidate below goes through at least autocontrast + unsharp
    # mask (_preprocess_for_ocr) — necessary rescue for a noisy real-world
    # scan, but confirmed to actively *hurt* an already-clean image: on a
    # sharp, high-contrast synthetic Armenian test page, the unenhanced
    # original round-tripped through Tesseract's "hye" model far more
    # accurately than any processed variant did (extra characters and
    # dropped words appeared only after processing). Armenian glyphs seem
    # more sensitive to this than Latin/Cyrillic, where the processed
    # candidates still won cleanly in the same test — so rather than
    # special-case the language, just always offer the genuinely
    # untouched original as one more candidate and let the existing
    # highest-score-wins comparison pick it when it's actually better.
    candidates = [
        (raw_img.convert("RGB"), 1.0, 1.0),
        (pil_pre, sx, sy),
        (ImageOps.invert(pil_pre), sx, sy),
        (binarized, sx, sy),
        (ImageOps.invert(binarized), sx, sy),
    ]
    candidates += [(v, sx, sy) for v in _adaptive_threshold_variants(pil_pre)]

    cv2_variant = _cv2_clahe_variant(raw_img)
    if cv2_variant is not None:
        candidates.append((cv2_variant, cv2_variant.width / raw_img.width, cv2_variant.height / raw_img.height))

    shading_variant = _shading_corrected_variant(raw_img)
    if shading_variant is not None:
        candidates.append((shading_variant, shading_variant.width / raw_img.width, shading_variant.height / raw_img.height))

    return candidates


def _detections_from_candidate(candidate: "Image.Image", sx: float, sy: float, lang: str, config: str, min_conf: float) -> List[tuple]:
    """Run OCR on one candidate rendering/config and return accepted
    (box, conf, word) detections in the original image's pixel coordinates."""
    try:
        data = pytesseract.image_to_data(candidate, lang=lang, config=config, output_type=pytesseract.Output.DICT)
    except Exception:
        return []
    out = []
    for i, word in enumerate(data.get("text", [])):
        word = word.strip()
        if not word or _REPEATED_CHAR_RE.search(word) or not _looks_like_text(word):
            continue
        try:
            conf = float(data["conf"][i])
        except (ValueError, TypeError):
            conf = -1.0
        if conf < min_conf:
            continue
        box = (
            data["left"][i] / sx,
            data["top"][i] / sy,
            data["width"][i] / sx,
            data["height"][i] / sy,
        )
        out.append((box, conf, word))
    return out


def _dedup_detections(detections: List[tuple]) -> List[tuple]:
    """Collapse duplicate/overlapping detections down to one (box, conf,
    word, support) per region, keeping the highest-confidence read and
    recording how many of the raw (pre-dedup) detections — i.e. how many
    different candidate renderings/configs — proposed something at
    roughly that same spot.

    That support count is what later tells genuine text apart from noise
    on busy illustrated images: real text tends to get found by several
    independently-thresholded renderings at roughly the same place, while
    a random patch of illustration texture only forms a plausible-looking
    "word" shape under one particular thresholding by chance, and rarely
    does so consistently across renderings — even though, taken alone,
    Tesseract can be just as confident about that one misread as it is
    about a real word.
    """
    all_boxes = [d[0] for d in detections]
    sorted_dets = sorted(detections, key=lambda d: -d[1])
    kept = []
    for box, conf, word in sorted_dets:
        x, y, w, h = box
        # A misread graphic (a ribbon edge, an arrow tip, a row of icons)
        # can get labeled as one "word" spanning a box far wider per
        # character than real text ever is — Tesseract's confidence
        # reflects how sure it is about the character shapes it guessed,
        # not whether the box is a plausible word at all.
        if w > len(word) * h * 1.3:
            continue
        if any(_boxes_compete(box, kb) for kb, _c, _w, _s in kept):
            continue
        support = sum(1 for b2 in all_boxes if _boxes_compete(box, b2))
        kept.append((box, conf, word, support))
    return kept


def _split_merged_word_detection(raw_img: "Image.Image", det: tuple, lang: str) -> List[tuple]:
    """If a single detected word's box has an internal ink gap wide enough
    to be a real inter-word space, split it into two detections there.

    Targets a failure mode the line-level space-*joining* elsewhere in
    this module can't reach: Tesseract's own word segmentation
    occasionally merges two visually-close words — tight kerning, an
    italic slant bridging the visual gap — into a single box to begin
    with, so there's no second detection to join a space onto; the fix
    has to come from the box's own pixels. Confirmed reproducible: two
    words rendered with a 1px gap in a bold italic font come back from
    Tesseract as one merged word ("Commentary"+"on" -> "Commentaryon"),
    while the same words at a 4px+ gap come back correctly split.

    Returns [det] unchanged whenever nothing meets the bar: word too
    short to bother, no numpy for the pixel analysis, no gap found, or a
    re-OCR of either half fails to produce plausible text — a failed
    split must never be worse than leaving the original merged word.
    """
    box, conf, word, support = det
    x, y, w, h = box
    if len(word) < 6 or h < 8 or np is None or pytesseract is None:
        return [det]

    crop = raw_img.crop((int(x), int(y), int(x + w), int(y + h))).convert("L")
    threshold = _otsu_threshold(crop)
    arr = np.array(crop)
    binarized = arr < threshold
    # Otsu doesn't know which side of the cutoff is ink vs. background;
    # assume ink is the minority class — true for any normal word crop,
    # where letters cover less area than the space around them.
    if binarized.sum() > binarized.size / 2:
        binarized = ~binarized

    col_ink = binarized.sum(axis=0)
    cw = col_ink.shape[0]
    margin = max(2, int(cw * 0.12))
    min_gap = max(3, int(h * 0.35))

    best_mid, best_width, run_start = None, 0, None
    for i in range(margin, cw - margin):
        if col_ink[i] == 0:
            if run_start is None:
                run_start = i
        elif run_start is not None:
            run_len = i - run_start
            if run_len >= min_gap and run_len > best_width:
                best_width, best_mid = run_len, (run_start + i) // 2
            run_start = None
    if run_start is not None:
        run_len = (cw - margin) - run_start
        if run_len >= min_gap and run_len > best_width:
            best_mid = (run_start + cw - margin) // 2

    if best_mid is None:
        return [det]

    left_crop = raw_img.crop((int(x), int(y), int(x + best_mid), int(y + h)))
    right_crop = raw_img.crop((int(x + best_mid), int(y), int(x + w), int(y + h)))
    try:
        left_text = pytesseract.image_to_string(left_crop, lang=lang, config="--oem 1 --psm 7").strip()
        right_text = pytesseract.image_to_string(right_crop, lang=lang, config="--oem 1 --psm 7").strip()
    except Exception:
        return [det]

    if not (left_text and right_text and _looks_like_text(left_text) and _looks_like_text(right_text)):
        return [det]
    if _REPEATED_CHAR_RE.search(left_text) or _REPEATED_CHAR_RE.search(right_text):
        return [det]

    return [
        ((x, y, best_mid, h), conf, left_text, support),
        ((x + best_mid, y, w - best_mid, h), conf, right_text, support),
    ]


_craft_model = None
_craft_model_tried = False


def _get_craft_model():
    """Lazily load and cache the optional CRAFT text-detection model.

    Same two-state caching as _get_doclayout_model, for the same reason:
    None is also the legitimate "unavailable" outcome, so a separate
    "tried" flag stops every image from retrying a slow/failed load.
    """
    global _craft_model, _craft_model_tried
    if _craft_model_tried:
        return _craft_model
    _craft_model_tried = True

    if _craft_load_model is None:
        return None
    try:
        weights_path = os.getenv("CRAFT_MODEL_PATH")
        if not weights_path:
            from huggingface_hub import hf_hub_download
            weights_path = hf_hub_download(
                repo_id="boomb0om/CRAFT-text-detector", filename="craft_mlt_25k.pth"
            )
        _craft_model = _craft_load_model(weights_path)
    except Exception:
        _craft_model = None
    return _craft_model


def _craft_line_detections(raw_img: "Image.Image", lang: str) -> List[tuple]:
    """Best-effort: detect text regions with CRAFT — a general-purpose
    text detector, independent of Tesseract's own layout analysis — and
    recognize each region with Tesseract, returning plain (box, conf,
    word) tuples in the same raw shape _detections_from_candidate already
    produces, so the caller can merge them straight into the same pool
    _dedup_detections runs over. That's deliberate: it means a CRAFT
    detection is only trusted as much as the existing, already-calibrated
    confidence/cross-detection-support logic already trusts anything else
    — a CRAFT-only read with nothing corroborating it gets exactly the
    same low support and the same strict trust bar as any other
    uncorroborated detection, rather than a new bespoke comparison. (Two
    earlier attempts at "let a second, independent OCR pass carry more
    weight" both caused real regressions on busy/illustrated pages by
    outvoting correct, conservative results — reusing the existing
    calibration instead of inventing a new one avoids repeating that.)

    Confirmed valuable on a real poster: Tesseract's own segmentation
    fragmented/dropped several comma-dense lines that CRAFT detects as one
    clean region each, read correctly once cropped tightly around just
    that region and recognized with --psm 7 (single line).

    Returns [] on any failure, missing dependency, or no boxes found.
    """
    net = _get_craft_model()
    if net is None or pytesseract is None:
        return []

    try:
        rgb = raw_img.convert("RGB")
        boxes = _craft_detect_boxes(net, rgb)
    except Exception:
        return []

    detections = []
    for x0, y0, x1, y1 in boxes:
        pad = 4
        crop = rgb.crop((max(0, x0 - pad), max(0, y0 - pad), x1 + pad, y1 + pad))
        try:
            text = pytesseract.image_to_string(crop, lang=lang, config="--oem 1 --psm 7").strip()
            data = pytesseract.image_to_data(crop, lang=lang, config="--oem 1 --psm 7", output_type=pytesseract.Output.DICT)
        except Exception:
            continue
        words = [w for w in data["text"] if w.strip()]
        if not text or not words or not all(_looks_like_text(w) for w in words) or _REPEATED_CHAR_RE.search(text):
            continue
        confs = [float(c) for c, w in zip(data["conf"], data["text"]) if w.strip() and float(c) >= 0]
        avg_conf = sum(confs) / len(confs) if confs else 0.0
        # No cross-rendering support count backs these up (see below), so
        # recognition confidence alone has to clear a real bar here.
        if avg_conf < 40:
            continue
        detections.append(((x0, y0, x1 - x0, y1 - y0), avg_conf, text))
    return detections


def _collect_ocr_detections(raw_img: "Image.Image", lang: str, min_conf: float) -> List[tuple]:
    """Run every candidate rendering × page-segmentation config and return
    every accepted (box, conf, word) detection, in raw_img's own pixel
    coordinates, with duplicates from overlapping candidates collapsed.

    A single "best overall candidate" doesn't exist for busy multi-color
    graphics: one thresholding pass may read a ribbon banner's yellow-on-
    red text well but miss a heading entirely, while another does the
    reverse. Collecting detections from every candidate and keeping the
    highest-confidence read of each *region* (rather than picking one
    candidate's full-image output) lets each region be read by whichever
    rendering actually suits it.
    """
    detections = []
    for candidate, sx, sy in _ocr_candidate_images(raw_img):
        for config in _OCR_CONFIGS:
            detections.extend(_detections_from_candidate(candidate, sx, sy, lang, config, min_conf))
    deduped = _dedup_detections(detections)
    split_out = []
    for det in deduped:
        split_out.extend(_split_merged_word_detection(raw_img, det, lang))
    return split_out


def _cluster_lines(detections: List[tuple]) -> List[dict]:
    """Group detections into horizontal lines by vertical overlap AND
    horizontal proximity to *some* other word already known to share the
    line — connected transitively (union-find over all pairs), not just
    "close to the nearest word processed so far".

    Vertical proximity alone isn't enough: a decorative glyph in a border
    far to the side of the page can share a similar y-coordinate with a
    real text line purely by coincidence and get merged into it.
    Requiring horizontal proximity keeps unrelated same-row content
    separate — but that check has to be transitive. A single greedy pass
    over detections sorted by y, matching each one against only the
    nearest word already accumulated in a line so far, is order-
    dependent: tiny per-word y jitter (ascenders, cap-height nudging one
    box a few pixels higher than its neighbors) can process a long
    line's words out of left-to-right order, splitting one real text
    line into two or three fragments before enough words have joined a
    chain to bridge them back together. Union-find over every pair
    doesn't have that ordering problem — a line forms correctly
    regardless of which order its words happen to be visited in.
    """
    n = len(detections)
    boxes = [d[0] for d in detections]
    parent = list(range(n))

    def find(i: int) -> int:
        while parent[i] != i:
            parent[i] = parent[parent[i]]
            i = parent[i]
        return i

    def union(i: int, j: int) -> None:
        ri, rj = find(i), find(j)
        if ri != rj:
            parent[ri] = rj

    for i in range(n):
        xi, yi, wi, hi = boxes[i]
        for j in range(i + 1, n):
            xj, yj, wj, hj = boxes[j]
            if abs(yi - yj) >= max(hi, hj) * 0.6:
                continue
            gap = max(xi - (xj + wj), xj - (xi + wi))
            if gap < max(hi, hj) * 8:
                union(i, j)

    groups: dict = {}
    for i, det in enumerate(detections):
        groups.setdefault(find(i), []).append(det)

    lines = []
    for dets in groups.values():
        y0 = min(d[0][1] for d in dets)
        y1 = max(d[0][1] + d[0][3] for d in dets)
        lines.append({"y": (y0 + y1) / 2, "dets": dets})
    lines.sort(key=lambda ln: ln["y"])
    return lines


def _detections_to_text(detections: List[tuple]) -> str:
    """Assemble kept detections into reading-order text: cluster into
    lines by vertical overlap, then sort each line left-to-right."""
    lines = _cluster_lines(detections)
    return "\n".join(
        " ".join(det[2] for det in sorted(ln["dets"], key=lambda d: d[0][0]))
        for ln in lines
    )


# Trust bar for a detection to stand on its own — outside a dominant
# paragraph block, or as part of what makes a candidate block "real" in
# the first place: both a high confidence AND a high cross-rendering
# support are required. Confidence alone doesn't discriminate — a random
# patch of illustration texture can score just as "confident" as real
# text under one particular thresholding — but real text is reliably
# found by many independently-thresholded renderings at the same spot,
# while that kind of noise rarely is. Calibrated against a busy
# illustrated poster (many high-confidence noise detections, support up
# to 13) and a set of clean chat-bubble captions (real short words with
# support as low as 4): this is the tightest bar that still eliminates
# 100% of the poster noise while keeping the large majority of genuine
# short captions/words.
_MIN_TRUSTED_SUPPORT = 8
_MIN_TRUSTED_CONFIDENCE = 85.0

# Below this many recognized words, the "trusted" full-page reading is
# sparse enough to call a real failure (Tesseract's single winning
# candidate lost most of the page), so a boxed-reconstruction fallback
# with more recovered words is worth trusting instead. At or above it,
# a word-count margin alone stopped being a safe signal: on a real dense
# scan, the union of detections across many renderings accumulates
# genuine per-word fragments (one real word split into 2-3 partial reads
# by different renderings) that inflate boxed_words exactly like a real
# recovery would, with no reliable way to tell the two apart from counts
# or confidence alone (both were tested and failed to discriminate).
_TRUSTED_SPARSE_WORD_LIMIT = 15


def _find_dominant_paragraph_block(detections: List[tuple]) -> Optional[dict]:
    """Find one clearly dominant multi-line paragraph block among the
    detections, if there is one, or None.

    Ornate decorative borders/frames (lace patterns, a ring of small
    icons around a greeting card) can generate dozens of scattered,
    medium-confidence word-shaped misreads that individually clear the
    normal bar. A real paragraph is reliably several consecutive,
    closely-spaced lines with multiple words each — border noise almost
    never clusters that way.
    """
    lines = _cluster_lines(detections)
    if len(lines) < 3:
        return None

    heights = sorted(d[0][3] for ln in lines for d in ln["dets"])
    median_h = heights[len(heights) // 2]
    if median_h <= 0:
        return None

    # Group lines into blocks by *both* vertical proximity and horizontal
    # overlap. Y-proximity alone isn't enough: a column of single-glyph
    # border noise running down one edge sits at roughly the same
    # y-positions as a paragraph's lines and would otherwise chain
    # together into a tall fake "block" purely from tight vertical
    # spacing, even though each of its "lines" has only one word. Also
    # requiring horizontal overlap keeps a side column of noise from ever
    # joining the paragraph's block, since they occupy disjoint x-ranges.
    blocks: List[dict] = []
    for ln in lines:
        x0 = min(d[0][0] for d in ln["dets"])
        x1 = max(d[0][0] + d[0][2] for d in ln["dets"])
        placed = False
        for blk in blocks:
            last = blk["lines"][-1]
            overlaps = x0 <= blk["x1"] and blk["x0"] <= x1
            if overlaps and (ln["y"] - last["y"]) < median_h * 1.8:
                blk["lines"].append(ln)
                blk["x0"], blk["x1"] = min(blk["x0"], x0), max(blk["x1"], x1)
                placed = True
                break
        if not placed:
            blocks.append({"lines": [ln], "x0": x0, "x1": x1})

    def word_count(blk):
        return sum(len(ln["dets"]) for ln in blk["lines"])

    def is_high_quality(blk) -> bool:
        # Density (lines × words/line) alone isn't a safe signal on a
        # heavily-noisy image: a busy illustration can produce so many
        # scattered low-quality detections that a "block" of pure noise
        # still clears the density bar by sheer volume. Requiring most of
        # a candidate block's own words to individually clear the same
        # trust bar used elsewhere (high confidence AND cross-rendering
        # support) is what actually tells a real paragraph — where nearly
        # every word is a solid, repeatedly-confirmed read — apart from
        # that kind of noise, where only a stray word or two would.
        dets = [d for ln in blk["lines"] for d in ln["dets"]]
        if not dets:
            return False
        trusted = sum(1 for d in dets if d[3] >= _MIN_TRUSTED_SUPPORT and d[1] >= _MIN_TRUSTED_CONFIDENCE)
        return trusted / len(dets) >= 0.6

    # A real paragraph has multiple words on nearly every line; scattered
    # border noise averages close to one word per "line" even when it
    # does chain together. Requiring a minimum density here is what
    # actually excludes it, not line count alone.
    paragraph_like = [
        b for b in blocks
        if len(b["lines"]) >= 3 and word_count(b) / len(b["lines"]) >= 3 and is_high_quality(b)
    ]
    if not paragraph_like:
        return None  # nothing clearly paragraph-shaped

    dominant = max(paragraph_like, key=lambda b: (len(b["lines"]), word_count(b)))
    if word_count(dominant) < 9:
        return None
    return dominant


def _prefer_dominant_paragraph_block(detections: List[tuple]) -> List[tuple]:
    """If the detections contain one clearly dominant paragraph block,
    trust it outright and apply the stricter confidence+support bar to
    everything outside it. With no dominant block (posters/captions made
    of a few separate short text elements, or a busy illustration with no
    real text structure at all), apply that same bar to everything.
    """
    dominant = _find_dominant_paragraph_block(detections)
    dominant_ids = {id(d) for ln in dominant["lines"] for d in ln["dets"]} if dominant else set()
    return [
        det for det in detections
        if id(det) in dominant_ids or (det[3] >= _MIN_TRUSTED_SUPPORT and det[1] >= _MIN_TRUSTED_CONFIDENCE)
    ]


def _full_text_score(text: str) -> int:
    return sum(len(w) for w in _WORD_RE.findall(text))


def _word_count(text: str) -> int:
    return len(_WORD_RE.findall(text))


# Full-page configs for the "trust Tesseract's own reading order" path —
# --psm 11 (sparse text, no particular order) is deliberately excluded
# here since it's the opposite of what a coherent document needs. --oem 1
# pinned for the same reason as _OCR_CONFIGS above.
_DOCUMENT_OCR_CONFIGS = ("--oem 1", "--oem 1 --psm 4", "--oem 1 --psm 6")


def _text_from_word_data(data: dict) -> str:
    """Reconstruct text from image_to_data's word-level output, preserving
    Tesseract's own block/paragraph/line structure — a faithful stand-in
    for image_to_string's serialization, not a different reading-order
    reconstruction, so the "trust Tesseract's own reading order" property
    _best_full_page_text relies on still holds.

    Words within a (block, par, line) group are joined sorted by their own
    x-position, not by Tesseract's raw word_num sequence — confirmed
    necessary on a real decorative, centered poster-style page, where
    word_num 1 within a claimed line sat *below* word_num 2-5 in the same
    line (different y entirely): Tesseract's own within-line ordering
    isn't reliably left-to-right on this kind of layout, even though its
    block/paragraph/line grouping and the order those groups appear in
    (top-to-bottom reading order) still are.
    """
    groups: dict = {}
    for i, word in enumerate(data["text"]):
        word = word.strip()
        if not word:
            continue
        key = (data["block_num"][i], data["par_num"][i], data["line_num"][i])
        groups.setdefault(key, []).append((data["left"][i], word))

    out_lines: List[str] = []
    prev_block_par = None
    for key, words in groups.items():
        if prev_block_par is not None and key[:2] != prev_block_par:
            out_lines.append("")
        words.sort(key=lambda w: w[0])
        out_lines.append(" ".join(w for _left, w in words))
        prev_block_par = key[:2]
    return "\n".join(out_lines)


def _mean_word_confidence(data: dict) -> float:
    confs = [float(c) for c, w in zip(data["conf"], data["text"]) if w.strip() and float(c) >= 0]
    return sum(confs) / len(confs) if confs else -1.0


# See _best_full_page_text's docstring for how these two work together:
# the confidence floor throws out pure hallucination before it can set
# the completeness reference, and the ratio then requires near-parity
# with the fullest plausible reading, not just "not obviously sparse".
_MIN_CANDIDATE_CONFIDENCE = 50.0
_FULL_PAGE_COMPLETENESS_RATIO = 0.85


def _best_full_page_text(raw_img: "Image.Image", lang: str) -> str:
    """Full-page OCR that trusts Tesseract's own built-in reading-order
    reconstruction, trying each candidate rendering/config and keeping
    whichever Tesseract is actually most confident about.

    Not the same as "recognized the most word-shaped characters" — a
    heavily processed candidate can hallucinate extra glyphs that still
    look like real words, inflating a plain word-character-count score,
    while scoring far lower on Tesseract's own per-character confidence
    (confirmed on a real case: a garbled candidate won the old
    character-count comparison outright while sitting ~20 points below
    every other candidate's mean confidence). Candidates that recognize
    much less text than a typical candidate are excluded first so a
    sparse-but-confident read (e.g. one easy word) can't beat a complete
    page merely by having fewer chances to be wrong.

    The completeness reference is the max word count *among candidates
    that clear a confidence floor first*, not the raw max and not the
    median. The raw max fails on a real scanned photo (paper
    texture/grain, not just clean synthetic text): one candidate can
    hallucinate a wall of spurious short "words" out of noise —
    confirmed on one where a single outlier candidate found 405 "words"
    at 28% confidence against every genuinely good candidate's 51-144 at
    56-92%. Using that outlier as the reference inflated the
    completeness bar so high it excluded every clean candidate. But the
    median has the opposite failure: on a real dense scan, Tesseract's
    automatic page-layout analysis can silently drop whole paragraphs
    from *most* candidates while still reading the paragraphs it keeps
    at high confidence (confirmed on one where several candidates
    independently settled on the same incomplete ~70-word reading at
    ~90% confidence, meaning the median sat right there too) — so a
    genuinely complete candidate at merely-good confidence (~75%, still
    clearing the floor) never got to compete, since it didn't clear
    "median * 0.7" either. Filtering to a confidence floor first (which
    the 405-word outlier fails outright) and taking the max *within that
    filtered pool* gets the reference from the actual fullest plausible
    reading instead of the typical one, so a candidate that's merely
    good but missing a third of the page can't beat it on confidence
    alone.
    """
    results = []
    for candidate, _sx, _sy in _ocr_candidate_images(raw_img):
        for config in _DOCUMENT_OCR_CONFIGS:
            try:
                data = pytesseract.image_to_data(candidate, lang=lang, config=config, output_type=pytesseract.Output.DICT)
            except Exception:
                continue
            word_count = sum(1 for w in data["text"] if w.strip())
            if word_count == 0:
                continue
            results.append((_mean_word_confidence(data), word_count, data))

    if not results:
        return ""

    plausible = [r for r in results if r[0] >= _MIN_CANDIDATE_CONFIDENCE]
    pool = plausible or results
    max_words = max(r[1] for r in pool)
    complete_enough = [r for r in pool if r[1] >= max_words * _FULL_PAGE_COMPLETENESS_RATIO]
    best_conf, _best_words, best_data = max(complete_enough, key=lambda r: r[0])
    return _text_from_word_data(best_data)


def _looks_like_dense_document(lines: List[dict]) -> bool:
    """Whether a probe pass's lines look like a normal multi-line text
    document (many lines with several words each) rather than scattered
    poster/graphic-style fragments.

    Deliberately not based on _find_dominant_paragraph_block's cross-line
    grouping: that function's "does line N join line N+1's block" gap
    threshold (median word-height × 1.8) was tuned against poster-style
    content, where word height and line spacing are both dictated by one
    big display font. Normal single-spaced body text has a much smaller
    word-bounding-box height relative to its line pitch (leading adds
    space a word's own ink never occupies), so that same threshold can
    fail to bridge even consecutive lines *within* one paragraph — this
    needs its own, simpler signal: just line/word density, no grouping.
    """
    dense_lines = [ln for ln in lines if len(ln["dets"]) >= 4]
    total_words = sum(len(ln["dets"]) for ln in lines)
    return len(dense_lines) >= 5 and total_words >= 20


_doclayout_model = None
_doclayout_model_tried = False


def _get_doclayout_model():
    """Lazily load and cache the optional DocLayout-YOLO model.

    Returns None — never raises — if the `doclayout-yolo` package isn't
    installed, its weights can't be obtained (no DOCLAYOUT_MODEL_PATH and
    no network access), or loading fails for any other reason. Cached in
    a *two-state* way (a "tried" flag alongside the value, not just
    checking `is None`) because None is also the legitimate "unavailable"
    outcome — without the separate flag, every image would retry the
    load (including a slow network timeout) rather than remembering it
    already failed once.
    """
    global _doclayout_model, _doclayout_model_tried
    if _doclayout_model_tried:
        return _doclayout_model
    _doclayout_model_tried = True

    if YOLOv10 is None:
        return None
    try:
        weights_path = os.getenv("DOCLAYOUT_MODEL_PATH")
        if not weights_path:
            # YOLOv10.from_pretrained() delegates to huggingface_hub's
            # PyTorchModelHubMixin, which expects a config.json describing
            # the checkpoint filename to fetch. This repo has none — just
            # the bare .pt file — so that auto-resolution silently falls
            # back to a hardcoded default ("yolov10n.pt") that doesn't
            # exist, raising FileNotFoundError (confirmed against both
            # huggingface_hub 1.x and 0.x). hf_hub_download() fetches the
            # exact named file directly, the same reliable, cached-in-
            # ~/.cache/huggingface mechanism the README already documents,
            # sidestepping the broken auto-resolution entirely.
            from huggingface_hub import hf_hub_download
            weights_path = hf_hub_download(
                repo_id="juliozhao/DocLayout-YOLO-DocStructBench",
                filename="doclayout_yolo_docstructbench_imgsz1024.pt",
            )
        _doclayout_model = YOLOv10(weights_path)
    except Exception:
        _doclayout_model = None
    return _doclayout_model


# Classes (of DocStructBench's 10) worth masking out before OCR: real
# illustration content that isn't text. Deliberately excludes
# figure_caption/table_caption/table_footnote (those ARE text — masking
# them would delete real content) and title/plain_text/abandoned_text
# (exactly what should reach OCR). isolated_formula is left unmasked too:
# formulas aren't a requested target and the false-masking risk isn't
# worth it for an unrequested class. "table" is deliberately excluded
# too, despite the class name: confirmed on a real bilingual vocabulary
# table that a detected "table" region is the table's *cell content*
# area, which routinely contains exactly the real text a table upload is
# for — masking it white before OCR destroyed the entire table's text
# outright, not just its gridlines/borders. Only genuine illustration
# content ("figure") has no text worth preserving.
_DOCLAYOUT_MASK_CLASSES = {"figure"}
_DOCLAYOUT_MASK_CONF = 0.25


# Below this, the correction is noise (sub-pixel jitter from the angle
# estimate itself), not real skew — skip the interpolation blur of rotating
# an already-straight page. Above this, the estimate is more likely a
# misdetection on non-text content (illustrations, borders) than real scan
# skew, which is rarely more than a few degrees — so it's left uncorrected
# rather than risk rotating a poster/graphic based on a bogus angle.
_DESKEW_MIN_ANGLE = 0.3
_DESKEW_MAX_ANGLE = 15.0


def _deskew(img: "Image.Image") -> "Image.Image":
    """Straighten a page that's rotated a few degrees off horizontal —
    common with phone photos of book pages and slightly crooked scans.

    Estimates the skew angle from the minimum-area bounding rectangle of
    all ink pixels (a page of text is, in aggregate, a long thin rotated
    rectangle) and rotates it back to horizontal. This corrects *rotation*
    only — not the 2D curl a page picks up near a book's spine, which needs
    a full page-dewarping model and is out of scope here; rotational skew
    is the far more common failure mode and the one cheap to fix locally.

    Returns img unchanged (never raises) on any failure, missing
    dependency, or when the estimated angle is outside the trusted range.
    """
    if cv2 is None or np is None:
        return img

    try:
        rgb = img.convert("RGB")
        arr = np.array(rgb)
        gray = cv2.cvtColor(arr, cv2.COLOR_RGB2GRAY)
        _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)

        coords = cv2.findNonZero(binary)
        if coords is None or len(coords) < 100:
            return img

        angle = cv2.minAreaRect(coords)[-1]
        # cv2.minAreaRect reports an angle in (-90, 0]; anything past -45
        # means it measured the rectangle's long side as vertical instead
        # of horizontal — rotate the other way to still land on horizontal.
        if angle < -45:
            angle = 90 + angle

        if abs(angle) < _DESKEW_MIN_ANGLE or abs(angle) > _DESKEW_MAX_ANGLE:
            return img

        h, w = gray.shape
        matrix = cv2.getRotationMatrix2D((w / 2, h / 2), angle, 1.0)
        rotated = cv2.warpAffine(arr, matrix, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)
        return Image.fromarray(rotated)
    except Exception:
        return img


def _mask_layout_noise(img: "Image.Image") -> "Image.Image":
    """Best-effort: if DocLayout-YOLO is available, detect figure/table
    regions and mask them white before OCR, so illustration noise can't
    confuse text detection — attacking the "busy image confuses OCR"
    problem at the source instead of filtering noise out after the fact.

    Returns img unchanged (never raises) on any failure, missing
    dependency, or when nothing meets the mask criteria.
    """
    model = _get_doclayout_model()
    if model is None or cv2 is None or np is None:
        return img

    try:
        rgb = img.convert("RGB")
        page_area = rgb.width * rgb.height
        results = model.predict(rgb, imgsz=1024, conf=_DOCLAYOUT_MASK_CONF, device="cpu", verbose=False)

        mask_candidates: List[tuple] = []
        text_boxes: List[tuple] = []
        for result in results:
            names = result.names
            for box in result.boxes:
                if float(box.conf[0]) < _DOCLAYOUT_MASK_CONF:
                    continue
                cls_name = names[int(box.cls[0])]
                coords = tuple(int(v) for v in box.xyxy[0].tolist())
                (mask_candidates if cls_name in _DOCLAYOUT_MASK_CLASSES else text_boxes).append(coords)

        def _overlaps_real_text(box: tuple) -> bool:
            x0, y0, x1, y1 = box
            box_area = max(1, (x1 - x0) * (y1 - y0))
            for tx0, ty0, tx1, ty1 in text_boxes:
                iw = max(0, min(x1, tx1) - max(x0, tx0))
                ih = max(0, min(y1, ty1) - max(y0, ty0))
                if (iw * ih) / box_area > 0.5:
                    return True
            return False

        boxes = []
        for x0, y0, x1, y1 in mask_candidates:
            # A genuine illustration is normally a discrete portion of
            # a page, not nearly the whole thing — confirmed on a
            # densely illuminated manuscript page, where the model's
            # own low-confidence "figure" detections span almost the
            # entire page (ornate borders + text intermixed confusing
            # the classifier) and masking them wiped out the one
            # actual text column along with the decoration, losing
            # everything instead of just the noise. A box this large
            # is more likely swallowing real content than isolating a
            # true illustration, so it's left unmasked.
            if (x1 - x0) * (y1 - y0) > page_area * 0.6:
                continue
            # The model can label the *same* region both a real text
            # class and "figure" at once — confirmed on a poster with a
            # decorative drop-shadow font on a solid color background,
            # which visually reads as graphic-like enough to get a
            # contradicting "figure" tag over the same spot as a "plain
            # text" one. Masking the figure tag there deleted the page's
            # only real text outright; a substantially-overlapping text
            # detection is direct evidence against the figure label.
            if _overlaps_real_text((x0, y0, x1, y1)):
                continue
            boxes.append((x0, y0, x1, y1))

        if not boxes:
            return img

        # Guard against death by a thousand cuts: several individually
        # under-60%-of-the-page boxes can still jointly cover nearly the
        # whole thing when they overlap or tile across it — confirmed on
        # the same manuscript page, where two ~45%-area "figure" boxes
        # together spanned the entire vertical extent, still wiping out
        # the one real text column even after the per-box filter above.
        mask = np.zeros(rgb.size[::-1], dtype=bool)
        for x0, y0, x1, y1 in boxes:
            mask[max(0, y0):y1, max(0, x0):x1] = True
        if mask.sum() > page_area * 0.6:
            return img

        arr = np.array(rgb)
        arr[mask] = 255
        return Image.fromarray(arr)
    except Exception:
        return img


def _recover_rotated_sidebar_text(raw_img: "Image.Image", covered_boxes: List[tuple], lang: str) -> List[tuple]:
    """Best-effort recovery of vertical/rotated text — a spine-style
    sidebar caption, a rotated stamp — that normal horizontal OCR doesn't
    just miss but actively misreads as scrambled/mirrored nonsense:
    recognizable letterforms in the wrong orientation still get "read",
    just garbled (confirmed reproducible: a vertical "PHILOSOPHIA ANTIQUA
    - VOLUME 137" sidebar comes back from horizontal OCR as fragments like
    "VIHdOSOTIHd" and "AINNIOA" — mirrored/reversed letter shapes, which
    without this function land as noise *inside* the main text, not just
    a gap where the sidebar should be) — so a plausibility filter on the
    normal pass alone won't catch it.

    Scans for narrow vertical bands with substantial ink that isn't
    explained by any already-detected text box, crops each, retries OCR
    at +90/-90 degrees, and keeps whichever rotation (if either) produces
    plausible multi-word text. Returns a list of (region_box, text) pairs
    — empty on any failure, missing dependency, or when no such band is
    found; this is a bonus recovery on top of the normal pipeline, never a
    requirement for it. The region box lets the caller also drop that
    area's original horizontal-orientation misreads out of the main
    detections, instead of just adding the recovered text alongside them.
    """
    if cv2 is None or np is None or pytesseract is None:
        return []
    try:
        rgb = raw_img.convert("RGB")
        gray = cv2.cvtColor(np.array(rgb), cv2.COLOR_RGB2GRAY)
        threshold = _otsu_threshold(Image.fromarray(gray))
        ink = (gray < threshold).astype("uint8")
        # Otsu doesn't know which side is ink; assume it's the minority
        # class over the whole page (true for any normal document/poster).
        if ink.sum() > ink.size / 2:
            ink = 1 - ink

        # Letters are thin, sparse strokes — even inside a solid block of
        # text, ink covers only a small fraction of the bounding region's
        # pixels. A vertical-biased dilation bridges those strokes into a
        # continuous column blob so a straightforward density scan can
        # find it; without it, real text's own per-column ink density is
        # too low to distinguish from noise at any usable threshold.
        dilated = cv2.dilate(ink, np.ones((21, 5), np.uint8), iterations=1).astype(bool)

        h, w = gray.shape
        covered = np.zeros((h, w), dtype="uint8")
        for (x, y, bw, bh) in covered_boxes:
            x0, y0 = max(0, int(x)), max(0, int(y))
            x1, y1 = min(w, int(x + bw)), min(h, int(y + bh))
            covered[y0:y1, x0:x1] = 1
        # Pad the coverage mask too, since the dilation above spreads a
        # detected word's own ink beyond its tight bounding box.
        covered = cv2.dilate(covered, np.ones((5, 5), np.uint8), iterations=1).astype(bool)

        uncovered = dilated & ~covered
        col_density = uncovered.sum(axis=0) / max(1, h)

        # A genuine vertical text band has ink sustained over a large
        # fraction of the page height, concentrated in a narrow column
        # range — unlike scattered decorative noise elsewhere on the page.
        candidate_cols = np.where(col_density > 0.08)[0]
        if candidate_cols.size == 0:
            return []

        # Merge nearby candidate columns into one run generously: a single
        # physical vertical text band can show small x-gaps between its
        # own rotated letters (individual glyph widths vary), and treating
        # those as separate regions would only recognize one fragment
        # while leaving the rest of the same band's horizontal-orientation
        # noise undetected and unexcluded from the main text.
        runs = []
        run_start = prev = candidate_cols[0]
        for c in candidate_cols[1:]:
            if c - prev > 25:
                runs.append((run_start, prev))
                run_start = c
            prev = c
        runs.append((run_start, prev))

        results = []
        for x0, x1 in runs:
            if not (10 <= x1 - x0 <= w * 0.25):
                continue
            rows = np.where(uncovered[:, x0:x1 + 1].any(axis=1))[0]
            if rows.size == 0 or rows.max() - rows.min() < h * 0.15:
                continue
            y0, y1 = int(rows.min()), int(rows.max())
            pad = 4
            region_box = (max(0, x0 - pad), max(0, y0 - pad), (x1 - x0) + 2 * pad, (y1 - y0) + 2 * pad)
            crop = rgb.crop((region_box[0], region_box[1], region_box[0] + region_box[2], region_box[1] + region_box[3]))

            # +90 vs -90 both recognize real letterforms — one direction
            # just reads them upside-down/mirrored, producing text that's
            # equally plausible by shape/length/word-count alone (see
            # "PHILOSOPHIA ANTIQUA" vs. its mirrored "VIHdOSOTIHd" reading,
            # both 5 space-separated tokens). Tesseract's own recognition
            # confidence is what actually tells them apart: it's
            # consistently much higher on the correctly-oriented reading.
            best_text, best_conf = None, 0.0
            for angle in (90, -90):
                try:
                    data = pytesseract.image_to_data(
                        crop.rotate(angle, expand=True), lang=lang, config="--oem 1 --psm 6",
                        output_type=pytesseract.Output.DICT,
                    )
                except Exception:
                    continue
                words = [w for w in data["text"] if w.strip()]
                text = " ".join(words)
                confs = [float(c) for c, w in zip(data["conf"], data["text"]) if w.strip() and float(c) >= 0]
                avg_conf = sum(confs) / len(confs) if confs else 0.0
                if (
                    len(words) >= 2 and not _REPEATED_CHAR_RE.search(text)
                    and avg_conf > best_conf and avg_conf >= _MIN_TRUSTED_CONFIDENCE
                ):
                    best_text, best_conf = text, avg_conf
            if best_text:
                results.append((region_box, best_text))
        return results
    except Exception:
        return []


def _exclude_sidebar_noise(detections: List[tuple], sidebar_results: List[tuple]) -> List[tuple]:
    """Drop detections that fall mostly inside a recovered rotated-sidebar
    region — see _recover_rotated_sidebar_text: those detections are the
    horizontal-orientation misreads of the same content, not separate real
    words, and would otherwise sit in the output alongside the correct
    recovered reading."""
    if not sidebar_results:
        return detections
    recovered_boxes = [rb for rb, _ in sidebar_results]
    return [d for d in detections if not any(_box_mostly_within(d[0], rb) for rb in recovered_boxes)]


def _append_sidebar_text(text: str, sidebar_results: List[tuple]) -> str:
    if not sidebar_results:
        return text
    return "\n\n".join([text] + [t for _, t in sidebar_results])


def _merge_craft_recovery(raw_img: "Image.Image", detections: List[tuple], lang: str) -> List[tuple]:
    """Merge in CRAFT-detected text regions (see _craft_line_detections),
    replacing any existing detection they cover rather than sitting
    alongside it — a small, wrong Tesseract fragment can otherwise end up
    right next to the correct, complete CRAFT reading of the very same
    spot in the final joined text.

    Applied after the dominant-block trust-bar filter runs, the same
    point sidebar recovery merges in — deliberately: CRAFT detections
    already passed their own two-stage check (CRAFT's own detection
    confidence, then a Tesseract-recognition-confidence floor in
    _craft_line_detections) and structurally can't accumulate the
    cross-rendering support count that trust bar requires, since CRAFT
    only ever produces one detection per region rather than several
    renderings' worth to agree with each other. Confirmed necessary: with
    CRAFT detections merged in *before* that filter instead, every one of
    them was rejected outright on support alone, even though for the
    poster this was built for, every single one was the fully correct
    reading and every rejected fragment was wrong.

    A CRAFT box is only accepted if it overlaps something Tesseract's own
    pipeline *also* detected at that position — confirmed necessary on an
    ornately illuminated manuscript page, where CRAFT correctly finds
    "text-shaped" regions inside the decorative border (flowers, birds,
    geometric knotwork all have character-like edges) that Tesseract's
    own detections never touched at all — CRAFT's own confidence on these
    is not a usable signal here (some scored higher than the poster's
    genuine content), so requiring Tesseract to have found *something*
    (however wrong) at the same spot first is the discriminator that
    actually works: it's what every accepted poster line had in common
    and every rejected manuscript false-positive lacked. The real cost is
    a CRAFT box Tesseract missed completely can't be recovered even where
    it would have been correct — an acceptable trade for not amplifying
    noise on the images most likely to have it.
    """
    craft_detections = _craft_line_detections(raw_img, lang)
    if not craft_detections:
        return detections
    corroborated = [
        (box, conf, word) for box, conf, word in craft_detections
        if any(_box_mostly_within(d[0], box) or _box_mostly_within(box, d[0]) for d in detections)
    ]
    if not corroborated:
        return detections
    craft_boxes = [d[0] for d in corroborated]
    kept = [d for d in detections if not any(_box_mostly_within(d[0], cb) for cb in craft_boxes)]
    kept.extend((box, conf, word, 0) for box, conf, word in corroborated)
    return kept


# Maps a detected dominant script to the OCR_LANGS component that reads
# it — only used when that component is actually part of the requested
# language set (see _detect_dominant_single_lang).
_SCRIPT_TO_LANG = {"armenian": "hye", "cyrillic": "rus", "latin": "eng"}


def _detect_dominant_single_lang(raw_img: "Image.Image", lang: str) -> Optional[str]:
    """If a page is overwhelmingly written in just one of the requested
    languages, return that language alone instead of the full combined
    string — running the combined multi-script model on a single-script
    page is what caused real, confirmed misreads: a short English phrase
    on a stylized cover came back as Cyrillic look-alike garbage under
    the combined "hye+rus+eng" model, but read correctly under plain
    "eng" alone (a short word elsewhere was similarly misread into
    Armenian look-alikes). Returns None for anything not clearly
    dominated by one script, so the combined model stays in play exactly
    where it's needed — a genuinely mixed-language page (e.g. a
    bilingual table) needs every requested language available at once.
    """
    components = lang.split("+")
    if len(components) < 2:
        return None  # only one language configured anyway

    try:
        primary = _preprocess_for_ocr(raw_img)
        sx, sy = primary.width / raw_img.width, primary.height / raw_img.height
        probe = _dedup_detections(_detections_from_candidate(primary, sx, sy, lang, "--oem 1", min_conf=40))
    except Exception:
        return None

    script_chars: Counter = Counter()
    total_chars = 0
    for _box, _conf, word, _support in probe:
        script = _dominant_script(word)
        if script:
            script_chars[script] += len(word)
        total_chars += len(word)
    if total_chars < 20:  # not enough signal on a near-empty probe to trust a switch
        return None

    top_script, top_count = script_chars.most_common(1)[0] if script_chars else (None, 0)
    if top_script is None or top_count / total_chars < 0.85:
        return None

    single_lang = _SCRIPT_TO_LANG.get(top_script)
    if single_lang is None or single_lang not in components:
        return None
    return single_lang


def _ocr_best_of(raw_img: "Image.Image", lang: str) -> str:
    single_lang = _detect_dominant_single_lang(raw_img, lang)
    if single_lang and single_lang != lang:
        return _ocr_best_of(raw_img, single_lang)

    raw_img = _mask_layout_noise(raw_img)
    # Cheap single-pass probe (one rendering, one config) to tell a normal
    # dense document apart from scattered poster/graphic-style text,
    # before paying for either strategy's full multi-candidate cost.
    primary = _preprocess_for_ocr(raw_img)
    sx, sy = primary.width / raw_img.width, primary.height / raw_img.height
    probe = _dedup_detections(_detections_from_candidate(primary, sx, sy, lang, "--oem 1", min_conf=40))
    if _looks_like_dense_document(_cluster_lines(probe)):
        # Trust Tesseract's own full-page reading-order reconstruction by
        # default: different renderings/configs yield slightly different
        # word boxes for the same text, and stitching several of them back
        # into lines can scramble reading order or drop a word that lost a
        # confidence tie-break — a real, observed failure mode on long
        # wrapped paragraphs.
        trusted_text = _best_full_page_text(raw_img, lang)

        # But cross-check it against the word-box reconstruction, which
        # structurally can't reproduce Tesseract's other common failure on
        # multi-column tables: image_to_string's own line/paragraph
        # serialization sometimes drops the space at a column or cell
        # boundary, gluing the last word of one cell to the first word of
        # the next (each word was individually recognized correctly — only
        # the space between them got lost). Joining each detected word with
        # an explicit space can't lose a space that way. If the box
        # reconstruction recovers meaningfully more words while covering at
        # least as much recognized text, the trusted path likely lost some
        # spaces — prefer the version that kept them.
        detections = _collect_ocr_detections(raw_img, lang, min_conf=40)
        # Check for a rotated sidebar/spine caption before the dominant-
        # block filter narrows things down — "already explained by a real
        # detection" should include everything found, not just what the
        # stricter dominant-block bar kept. Note this only cleans the
        # *boxed* reconstruction below; if `trusted_text` (Tesseract's own
        # full-page string, with no per-word boxes to filter) wins the
        # comparison instead, any inline sidebar misreads there aren't
        # removed — only the recovered text gets appended either way.
        sidebar_results = _recover_rotated_sidebar_text(raw_img, [d[0] for d in detections], lang)
        detections = _exclude_sidebar_noise(detections, sidebar_results)
        detections = _prefer_dominant_paragraph_block(detections)
        boxed_text = _detections_to_text(detections)

        trusted_words, trusted_chars = _word_count(trusted_text), _full_text_score(trusted_text)
        boxed_words, boxed_chars = _word_count(boxed_text), _full_text_score(boxed_text)
        trusted_is_sparse = trusted_words < _TRUSTED_SPARSE_WORD_LIMIT
        text = (
            boxed_text
            if (trusted_is_sparse and boxed_words > trusted_words * 1.15 and boxed_chars >= trusted_chars * 0.9)
            else trusted_text
        )
        text = _strip_stray_script_glyphs(text)
        return _append_sidebar_text(text, sidebar_results)

    detections = _collect_ocr_detections(raw_img, lang, min_conf=40)
    sidebar_results = _recover_rotated_sidebar_text(raw_img, [d[0] for d in detections], lang)
    detections = _exclude_sidebar_noise(detections, sidebar_results)
    detections = _prefer_dominant_paragraph_block(detections)
    # CRAFT (an independent, non-Tesseract text detector) only for the
    # scattered/poster-style branch, not dense documents: this is where it
    # was validated (Tesseract's own segmentation fragmenting/dropping
    # comma-dense lines on a real poster) and it costs a real detection +
    # per-region OCR pass, not worth paying on every ordinary page without
    # the same proof it helps there too.
    detections = _merge_craft_recovery(raw_img, detections, lang)
    text = _strip_stray_script_glyphs(_detections_to_text(detections))
    return _append_sidebar_text(text, sidebar_results)


def extract_text_from_pdf(path: Path, lang: str = DEFAULT_OCR_LANGS) -> str:
    # First try to extract text directly (for digitally generated PDFs)
    if pdf_extract_text is not None:
        try:
            text = pdf_extract_text(str(path))
            if text and text.strip():
                return text
        except Exception:
            pass

    # Fallback: render pages to images and OCR
    if convert_from_path is None or pytesseract is None or Image is None:
        raise RuntimeError("Missing pdf->image OCR dependencies: install pdf2image, pytesseract, pillow")

    lang = _resolve_ocr_langs(lang)
    pages = convert_from_path(str(path), poppler_path=POPPLER_PATH)
    out_lines = []
    for i, page in enumerate(pages, start=1):
        # Route through the same pipeline direct image uploads use
        # (multi-candidate rendering/config search, dense-vs-sparse
        # routing, stray-script-glyph filtering, in-box merged-word
        # splitting, rotated-sidebar recovery) rather than one plain OCR
        # pass — a scanned page can have exactly the same table/cover/
        # sidebar layouts a standalone image upload does, and there's no
        # reason those fixes should only apply to one of this app's two
        # OCR entry points. _mask_layout_noise runs inside _ocr_best_of;
        # only deskew needs to happen first, here.
        page = _deskew(page)
        txt = _ocr_best_of(page, lang)
        out_lines.append(f"\n--- PAGE {i} ---\n")
        out_lines.append(txt)
    return "\n".join(out_lines)


def extract_text_from_image(path: Path, lang: str = DEFAULT_OCR_LANGS) -> str:
    if pytesseract is None or Image is None:
        raise RuntimeError("Missing image OCR dependencies: install pytesseract and pillow")
    # Closing the handle (rather than leaving it open until GC) matters on
    # Windows, where the caller's temp-file cleanup would otherwise fail
    # with "file in use by another process".
    lang = _resolve_ocr_langs(lang)
    with Image.open(path) as img:
        img = _deskew(img)
        return _ocr_best_of(img, lang)


def normalize_dark_image_to_paper(img: "Image.Image") -> Optional["Image.Image"]:
    """If the image looks like light text on a dark background, convert it
    to a normal-looking scanned-paper style: white background, black text.

    Dark-mode screenshots and dark poster/flyer designs are readable to a
    person but visually the inverse of a normal document; producing a
    black-on-white version alongside the extracted text gives something
    that looks like what people expect a "cleaned up" text image to look
    like. Returns None if the image isn't dark-background to begin with
    (nothing to normalize) or if Pillow isn't installed.
    """
    if Image is None or ImageOps is None or ImageStat is None:
        return None

    gray = ImageOps.grayscale(img.convert("RGB"))
    if ImageStat.Stat(gray).mean[0] >= 127:
        return None  # already a light background

    threshold = _otsu_threshold(gray)
    binarized = gray.point(lambda p: 255 if p > threshold else 0)

    # Otsu doesn't know which side of the cutoff is "background" vs
    # "text" — assume whichever value covers more pixels is the
    # background, and normalize so background is always white (255) and
    # text is always black (0), regardless of the source's original
    # polarity or color.
    hist = binarized.histogram()
    if hist[0] > hist[255]:
        binarized = ImageOps.invert(binarized)

    return binarized.convert("RGB")


def extract_text_from_textfile(path: Path) -> str:
    # Text/CSV files already are text — read them as-is rather than OCR-ing.
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_bytes().decode("utf-8", errors="replace")


def extract_text_from_spreadsheet(path: Path) -> str:
    if openpyxl is None:
        raise RuntimeError("Missing spreadsheet dependency: install openpyxl")
    wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    out_lines = []
    try:
        for ws in wb.worksheets:
            out_lines.append(f"\n--- SHEET: {ws.title} ---\n")
            for row in ws.iter_rows(values_only=True):
                cells = ["" if c is None else str(c) for c in row]
                if any(cells):
                    out_lines.append("\t".join(cells))
    finally:
        wb.close()
    return "\n".join(out_lines)


def extract_text_from_docx(path: Path) -> str:
    if python_docx is None:
        raise RuntimeError("Missing docx dependency: install python-docx")
    document = python_docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _strip_text_from_image(img: "Image.Image", lang: str = DEFAULT_OCR_LANGS) -> "Image.Image":
    """Remove OCR-detected text from an image via inpainting, leaving just
    the underlying artwork/photo.

    "Extract images" is for pulling out the picture asset itself — a
    poster's illustration, a product photo — not a duplicate of what
    "extract text" already returns baked into the picture. Reuses the
    same multi-candidate detection battery as text extraction (plain
    single-pass detection reliably misses decorative/stylized fonts on
    busy backgrounds, leaving the original text untouched) but with a
    much lower confidence floor: finding *where* text roughly is only
    needs to be right enough to mask it, not right enough to read it, so
    it's fine to be far more liberal here than when extracting text.
    Falls back to returning the image unchanged if OpenCV isn't
    installed, the language pack is missing, or OCR/inpainting fails for
    any reason — stripping is a best-effort enhancement, not something
    that should block image extraction.
    """
    if cv2 is None or np is None or pytesseract is None:
        return img

    try:
        resolved_lang = _resolve_ocr_langs(lang)
    except RuntimeError:
        return img

    rgb = img.convert("RGB")
    try:
        # A near-zero floor (below ~15) starts admitting misreads of
        # repeating decorative patterns (a row of pine-tree/dot/snowflake
        # shapes) as one long, low-confidence "word" spanning the whole
        # row — genuine text, even faint or stylized, reliably scores
        # well above that.
        detections = _collect_ocr_detections(rgb, resolved_lang, min_conf=15)
    except Exception:
        return img

    if not detections:
        return img

    mask = np.zeros((rgb.height, rgb.width), dtype=np.uint8)
    for det in detections:
        x, y, w, h = det[0]
        pad = max(3, int(0.3 * h))
        x0, y0 = max(0, int(x - pad)), max(0, int(y - pad))
        x1, y1 = min(rgb.width, int(x + w + pad)), min(rgb.height, int(y + h + pad))
        mask[y0:y1, x0:x1] = 255

    bgr = cv2.cvtColor(np.array(rgb), cv2.COLOR_RGB2BGR)
    inpainted = cv2.inpaint(bgr, mask, inpaintRadius=7, flags=cv2.INPAINT_TELEA)
    return Image.fromarray(cv2.cvtColor(inpainted, cv2.COLOR_BGR2RGB))


def _to_png_bytes(img: "Image.Image") -> bytes:
    if img.mode not in ("RGB", "RGBA", "L"):
        img = img.convert("RGBA" if "A" in img.getbands() else "RGB")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def extract_images_from_pdf(path: Path) -> List[bytes]:
    """Return every embedded image in the PDF as PNG bytes, in document order.

    Falls back to rendering full pages when the PDF has no embedded images
    (e.g. it's a scanned PDF stored as page-level vector/XObject content).
    """
    if pymupdf is None:
        raise RuntimeError("Missing PDF image extraction dependency: install pymupdf")

    doc = pymupdf.open(str(path))
    images: List[bytes] = []
    try:
        for page in doc:
            for img_info in page.get_images(full=True):
                xref = img_info[0]
                try:
                    extracted = doc.extract_image(xref)
                    raw = extracted.get("image")
                    if not raw:
                        continue
                    if Image is not None:
                        images.append(_to_png_bytes(_strip_text_from_image(Image.open(io.BytesIO(raw)))))
                    else:
                        images.append(raw)
                except Exception:
                    continue

        if not images and convert_from_path is not None:
            for page in convert_from_path(str(path), poppler_path=POPPLER_PATH):
                images.append(_to_png_bytes(page))
    finally:
        doc.close()

    return images


def extract_images_from_office(path: Path) -> List[bytes]:
    """xlsx/docx are zip archives; pull out whatever pictures are embedded
    under their media/ folder."""
    images: List[bytes] = []
    with zipfile.ZipFile(path) as z:
        for name in sorted(n for n in z.namelist() if "/media/" in n):
            raw = z.read(name)
            if Image is not None:
                try:
                    images.append(_to_png_bytes(_strip_text_from_image(Image.open(io.BytesIO(raw)))))
                    continue
                except Exception:
                    pass
            images.append(raw)
    return images


def extract_images_from_file(path: Path) -> List[bytes]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_images_from_pdf(path)
    if suffix in IMAGE_EXTS:
        if Image is None:
            raise RuntimeError("Missing image dependency: install pillow")
        with Image.open(path) as img:
            return [_to_png_bytes(_strip_text_from_image(img))]
    if suffix in OFFICE_ZIP_EXTS:
        return extract_images_from_office(path)
    if suffix in TEXT_EXTS:
        return []
    raise ValueError(f"Unsupported file type: {suffix}")


ALL_SUPPORTED_EXTS = IMAGE_EXTS | TEXT_EXTS | SPREADSHEET_EXTS | DOCX_EXTS | {".pdf"}


def extract_text_from_path(p: Path) -> Optional[str]:
    suffix = p.suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(p)
    if suffix in IMAGE_EXTS:
        return extract_text_from_image(p)
    if suffix in TEXT_EXTS:
        return extract_text_from_textfile(p)
    if suffix in SPREADSHEET_EXTS:
        return extract_text_from_spreadsheet(p)
    if suffix in DOCX_EXTS:
        return extract_text_from_docx(p)
    return None


def process_path(p: Path) -> Optional[str]:
    if not p.exists():
        print(f"Not found: {p}", file=sys.stderr)
        return None

    if p.is_dir():
        outputs = []
        for f in sorted(p.rglob("*")):
            if f.is_file() and f.suffix.lower() in ALL_SUPPORTED_EXTS:
                txt = process_path(f)
                if txt:
                    outputs.append(f"\n===== {f} =====\n")
                    outputs.append(txt)
        return "\n".join(outputs)

    # single file
    if p.suffix.lower() in ALL_SUPPORTED_EXTS:
        return extract_text_from_path(p)

    print(f"Unsupported file type: {p}", file=sys.stderr)
    return None


def main():
    parser = argparse.ArgumentParser(description="Extract text from PDFs and images (with OCR fallback).")
    parser.add_argument("path", help="File or directory path to process")
    parser.add_argument("-o", "--out", help="Output file. If omitted prints to stdout")
    args = parser.parse_args()

    p = Path(args.path)
    try:
        result = process_path(p)
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(2)

    if not result:
        print("No text extracted.")
        sys.exit(0)

    if args.out:
        outp = Path(args.out)
        outp.write_text(result, encoding="utf-8")
        print(f"Wrote: {outp}")
    else:
        print(result)


if __name__ == "__main__":
    main()
